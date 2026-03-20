"""Generate test fixtures programmatically — run once with: uv run python tests/create_fixtures.py"""
import pathlib

FIXTURES = pathlib.Path(__file__).parent / "fixtures"


def make_docx():
    from docx import Document
    from docx.shared import Pt
    doc = Document()

    doc.add_heading("Research Overview", level=1)
    doc.add_heading("Background", level=2)
    p = doc.add_paragraph()
    run = p.add_run("This is ")
    run2 = p.add_run("bold text")
    run2.bold = True
    run3 = p.add_run(" and ")
    run4 = p.add_run("italic text")
    run4.italic = True
    run5 = p.add_run(".")

    doc.add_heading("Data Table", level=2)
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    headers = ["Name", "Score", "Grade"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [("Alice", "95", "A"), ("Bob", "82", "B")]
    for r, row in enumerate(data, start=1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = val

    doc.add_heading("Nested List", level=2)
    doc.add_paragraph("Item one", style="List Bullet")
    doc.add_paragraph("Item two", style="List Bullet")

    doc.save(FIXTURES / "sample.docx")
    print("Created sample.docx")


def make_xlsx():
    from openpyxl import Workbook
    from openpyxl.styles import Font

    wb = Workbook()
    ws = wb.active
    ws.title = "Sales"

    # Bold header row
    headers = ["Region", "Q1", "Q2", "Q3", "Q4"]
    for col, h in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = Font(bold=True)

    # Merged cell spanning two columns
    ws.merge_cells("A2:A3")
    ws["A2"] = "North"
    data = [
        [None, 100, 120, 130, 110],
        [None, 90,  95,  105, 100],
        ["South", 200, 210, 190, 220],
    ]
    for r, row in enumerate(data, start=2):
        for c, val in enumerate(row, start=1):
            if val is not None:
                ws.cell(row=r, column=c, value=val)

    wb.save(FIXTURES / "sample.xlsx")
    print("Created sample.xlsx")


def make_pdf():
    import pymupdf  # type: ignore
    doc = pymupdf.open()
    page = doc.new_page()
    # Title — large font
    page.insert_text((72, 80), "Research Findings", fontsize=24, color=(0, 0, 0))
    # Section heading
    page.insert_text((72, 130), "Introduction", fontsize=16, color=(0, 0, 0))
    # Body text
    page.insert_text(
        (72, 160),
        "This document summarises the key findings of the study.",
        fontsize=11,
    )
    page.insert_text((72, 200), "Methods", fontsize=16, color=(0, 0, 0))
    page.insert_text((72, 230), "Data was collected over six months.", fontsize=11)
    doc.save(FIXTURES / "sample.pdf")
    print("Created sample.pdf")


if __name__ == "__main__":
    FIXTURES.mkdir(exist_ok=True)
    make_docx()
    make_xlsx()
    make_pdf()
    print("All fixtures created.")
